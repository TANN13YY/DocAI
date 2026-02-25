import os
import io
import asyncio
import uuid
import time
import re
import json
from contextlib import asynccontextmanager
from typing import List, Optional, Dict, Any

from fastapi import FastAPI, UploadFile, File, HTTPException, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import pdfplumber
import docx
from google import genai
from dotenv import load_dotenv

# Import database functions
from database import (
    init_db, 
    add_review, 
    get_reviews, 
    save_summary, 
    get_summary, 
    init_contact_table, 
    add_contact_submission,
    migrate_contact_table
)

load_dotenv()

# Global variables
client: Optional[genai.Client] = None
documents: Dict[str, str] = {} # In-memory document store

@asynccontextmanager
async def lifespan(app: FastAPI):
    # Startup
    global client
    api_key = os.getenv("GEMINI_API_KEY")
    if api_key:
        client = genai.Client(api_key=api_key)
        print("Gemini Client initialized.")
    else:
        print("Warning: GEMINI_API_KEY not found in environment variables.")
    
    init_db()
    init_contact_table()
    migrate_contact_table() # Ensure migration runs
    print("Database initialized.")
    
    yield
    
    # Shutdown
    print("Shutting down...")

app = FastAPI(lifespan=lifespan)

# Configure CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# --- Helper Functions ---

async def get_gemini_response_async(text: str) -> str:
    if not client:
        raise HTTPException(status_code=500, detail="Gemini API Key not configured.")
    
    try:
        prompt = f"""
        You are an expert academic tutor. Create a comprehensive, high-quality study guide for the following text.
        
        **CRITICAL INSTRUCTIONS**:
        1. **FULL COVERAGE**: You must analyze the **ENTIRE** provided text.
        2. **STRICT FORMATTING**: 
           - **DO NOT indent headings**.
           - Do not use code blocks for normal text.
        
        The study guide MUST include:
        1. **Executive Summary**: Comprehensive overview.
        2. **Key Concepts**: Structured list with definitions.
        3. **Study Notes**: Detailed explanations.
        4. **Practice Questions**: 10 distinct multiple-choice questions.
        
        Format the output in clean, professional Markdown. 
        
        Text to process:
        {text[:250000]}
        """
        
        # Retry logic
        max_retries = 3
        for attempt in range(max_retries):
            try:
                # Use async call
                response = await client.aio.models.generate_content(
                    model='gemini-2.5-flash',
                    contents=prompt
                )
                
                cleaned_text = response.text
                # Regex to un-indent headers
                cleaned_text = re.sub(r'^\s+(#+)', r'\1', cleaned_text, flags=re.MULTILINE)
                return cleaned_text
                
            except Exception as e:
                if "429" in str(e) and attempt < max_retries - 1:
                    wait_time = 2 ** attempt
                    print(f"Quota exceeded, retrying in {wait_time} seconds...")
                    await asyncio.sleep(wait_time)
                else:
                    raise e
                    
    except Exception as e:
        print(f"Error generating content: {e}")
        raise HTTPException(status_code=500, detail=f"AI generation failed: {str(e)}")

def process_pdf_sync(content: bytes) -> str:
    text = ""
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        total_pages = len(pdf.pages)
        print(f"PDF has {total_pages} pages.")
        for i, page in enumerate(pdf.pages):
            extracted = page.extract_text()
            if extracted:
                text += extracted + "\n"
    return text

def process_docx_sync(content: bytes) -> str:
    text = ""
    doc = docx.Document(io.BytesIO(content))
    print(f"Word doc has {len(doc.paragraphs)} paragraphs.")
    for para in doc.paragraphs:
        text += para.text + "\n"
    return text

# --- Models ---

class ReviewRequest(BaseModel):
    name: str
    role: Optional[str] = None
    content: str
    rating: int

class ShareRequest(BaseModel):
    content: str

class ContactRequest(BaseModel):
    name: str
    email: str
    subject: str
    description: str
    issue_resolved: int = 0

class TranslationRequest(BaseModel):
    text: str
    target_language: str

class ChatRequest(BaseModel):
    doc_id: str
    messages: list
    question: str

# --- Endpoints ---

@app.get("/")
def read_root():
    return {"message": "PDF Study Summarizer API is running"}

@app.get("/reviews")
def read_reviews():
    return get_reviews()

@app.post("/reviews")
def create_review(review: ReviewRequest):
    add_review(review.name, review.role, review.content, review.rating)
    return {"message": "Review added successfully"}

@app.post("/share")
def create_share_link(request: ShareRequest):
    share_id = save_summary(request.content)
    return {"share_id": share_id}

@app.get("/share/{share_id}")
def read_shared_summary(share_id: str):
    content = get_summary(share_id)
    if not content:
        raise HTTPException(status_code=404, detail="Summary not found")
    return {"content": content}

@app.post("/contact")
async def contact_form(request: ContactRequest, background_tasks: BackgroundTasks):
    submission_id = str(uuid.uuid4())
    timestamp = str(time.strftime('%Y-%m-%d %H:%M:%S'))
    
    try:
        # Check if add_contact_submission is sync (it is), so we can call it directly.
        # SQLite writes are fast enough to keep sync for now, or could offload.
        add_contact_submission(
            id=submission_id,
            name=request.name,
            email=request.email,
            subject=request.subject,
            description=request.description,
            timestamp=timestamp,
            issue_resolved=request.issue_resolved
        )
        

        
        return {"status": "success", "message": "We will get in touch soon."}
    except Exception as e:
        print(f"Database error: {e}")
        raise HTTPException(status_code=500, detail="Internal server error")

@app.post("/translate")
async def translate_text(request: TranslationRequest):
    print(f"Received translation request to {request.target_language}")
    if not client:
        raise HTTPException(status_code=500, detail="Gemini API Key not configured.")
    
    try:
        prompt = f"""
        Translate the following academic study guide into {request.target_language}.
        
        Maintain the original Markdown formatting.
        Do not translate Mermaid.js code blocks.
        
        Text to translate:
        {request.text}
        """
        
        # Retry logic
        max_retries = 3
        for attempt in range(max_retries):
            try:
                response = await client.aio.models.generate_content(
                    model='gemini-2.5-flash', # Using flash model for speed
                    contents=prompt
                )
                return {"translated_text": response.text}
            except Exception as e:
                if "429" in str(e) and attempt < max_retries - 1:
                    wait_time = 2 ** attempt
                    print(f"Quota exceeded, retrying in {wait_time} seconds...")
                    await asyncio.sleep(wait_time)
                else:
                    raise e
    except Exception as e:
        print(f"Translation error: {e}")
        raise HTTPException(status_code=500, detail=f"Translation failed: {str(e)}")

@app.post("/upload")
async def upload_file(file: UploadFile = File(...)):
    print(f"Received upload request: {file.filename}")
    
    if not file.filename.endswith((".pdf", ".docx")):
        raise HTTPException(status_code=400, detail="File must be a PDF or Word (.docx) document")

    try:
        # Read file content - awaiting file.read() is non-blocking in FastAPI
        print(f"Reading file content...")
        content = await file.read()
        text = ""

        # Offload CPU-bound processing to thread header
        if file.filename.lower().endswith(".pdf"):
            print(f"Processing PDF in thread...")
            text = await asyncio.to_thread(process_pdf_sync, content)
        
        elif file.filename.lower().endswith(".docx"):
            print(f"Processing Word Document in thread...")
            text = await asyncio.to_thread(process_docx_sync, content)
        
        print(f"Text extraction complete. Length: {len(text)} characters.")
        if not text.strip():
            print("Extraction failed (empty text).")
            raise HTTPException(status_code=400, detail="Could not extract text. It might be scanned or empty.")

        doc_id = str(uuid.uuid4())
        documents[doc_id] = text
        print(f"Stored document text with ID: {doc_id}")

        print(f"Calling Gemini API asynchronously...")
        study_guide = await get_gemini_response_async(text)
        print("Gemini response received.")
        
        return {"filename": file.filename, "study_guide": study_guide, "doc_id": doc_id}

    except HTTPException as he:
        raise he
    except Exception as e:
        print(f"Error in upload: {e}")
        raise HTTPException(status_code=500, detail=f"An error occurred: {str(e)}")

@app.post("/chat")
async def chat_with_document(request: ChatRequest):
    if request.doc_id not in documents:
        raise HTTPException(status_code=404, detail="Document context not found. Please re-upload.")
    
    doc_text = documents[request.doc_id]
    
    if not client:
         raise HTTPException(status_code=500, detail="Gemini API Key not configured.")

    try:
        # Construct chat prompt
        chat_prompt = f"""
        You are a helpful AI tutor assistant.
        
        Document Content (Truncated):
        {doc_text[:100000]}...
        
        Chat History:
        {request.messages}
        
        User Question: {request.question}
        
        Answer concise and helpful:
        """
        max_retries = 3
        for attempt in range(max_retries):
            try:
                response = await client.aio.models.generate_content(
                    model='gemini-2.5-flash',
                    contents=chat_prompt
                )
                return {"answer": response.text}
            except Exception as e:
                if "429" in str(e) and attempt < max_retries - 1:
                    wait_time = 2 ** attempt
                    print(f"Quota exceeded in chat, retrying in {wait_time} seconds...")
                    await asyncio.sleep(wait_time)
                else:
                    raise e
    except Exception as e:
        print(f"Chat error: {e}")
        return {"answer": "I'm sorry, I encountered an error while processing your question via AI."}
class QuizRequest(BaseModel):
    doc_id: Optional[str] = None
    text: Optional[str] = None

@app.post("/quiz")
async def generate_quiz(request: QuizRequest):
    doc_text = ""
    if request.doc_id and request.doc_id in documents:
        doc_text = documents[request.doc_id]
    elif request.text:
        doc_text = request.text
    else:
        raise HTTPException(status_code=400, detail="Either doc_id or text must be provided")
    
    prompt = f"""Based on the following text, generate a quiz with 10 multiple-choice questions.
    Return a JSON array of objects, where each object has:
    - "question": The question string
    - "options": A list of 4 answer options (strings)
    - "correct_answer": The index of the correct answer (0-3)

    Text (Truncated):
    {doc_text[:10000]}
    """ # Limiting text for context window

    try:
        if not client:
             raise HTTPException(status_code=500, detail="Gemini client not initialized")

        max_retries = 3
        for attempt in range(max_retries):
            try:
                response = await client.aio.models.generate_content(
                    model="gemini-2.5-flash",
                    contents=prompt,
                    config={
                        'response_mime_type': 'application/json'
                    }
                )
                return json.loads(response.text)
            except Exception as e:
                if "429" in str(e) and attempt < max_retries - 1:
                    wait_time = 2 ** attempt
                    print(f"Quota exceeded in quiz, retrying in {wait_time} seconds...")
                    await asyncio.sleep(wait_time)
                else:
                    raise e
    except Exception as e:
        print(f"Quiz generation error: {e}")
        raise HTTPException(status_code=500, detail=str(e))
